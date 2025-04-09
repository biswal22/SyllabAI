from flask import Flask, request, jsonify
from flask_cors import CORS
import pytesseract
from pdf2image import convert_from_path
from docx import Document
import tempfile
import os
from PIL import Image
import logging
import fitz  # PyMuPDF for better PDF handling
import sys
from openai import OpenAI
from dotenv import load_dotenv
import asyncio
import json
import re
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address

# Load environment variables
load_dotenv()

app = Flask(__name__)

# Configure CORS for different environments
if os.environ.get('FLASK_ENV') == 'production':
    # In production, only allow requests from the frontend domain
    frontend_url = os.environ.get('FRONTEND_URL', '*')
    CORS(app, resources={r"/*": {"origins": frontend_url}})
    print(f"Running in production mode, CORS configured for: {frontend_url}")
else:
    # In development, allow all origins
    CORS(app)
    print("Running in development mode, CORS configured for all origins")

# Initialize rate limiter with specific limits per user IP
limiter = Limiter(
    get_remote_address,
    app=app,
    default_limits=["24 per day", "12 per hour"],
    storage_uri=os.environ.get('REDIS_URL', 'memory://'),
)

# Set up logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

# Initialize OpenAI client
client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

def extract_text_from_pdf(file_path):
    logger.debug(f"Attempting to extract text from PDF: {file_path}")
    
    try:
        # First attempt: Try PyMuPDF (usually better for searchable PDFs)
        logger.debug("Attempting PyMuPDF extraction...")
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text() + "\n"
        doc.close()
        
        if text.strip():
            logger.debug("Successfully extracted text using PyMuPDF")
            return text
        
        # If no text found, try OCR
        logger.debug("No text found with PyMuPDF, attempting OCR...")
        
        images = convert_from_path(file_path)
        text = ""
        
        for i, image in enumerate(images):
            logger.debug(f"Processing page {i+1} with OCR")
            page_text = pytesseract.image_to_string(image)
            text += page_text + "\n"
            
        if not text.strip():
            raise Exception("No text extracted from PDF")
            
        return text
            
    except Exception as e:
        logger.error(f"Error in PDF extraction: {str(e)}")
        raise

def extract_text_from_docx(file_path):
    logger.debug(f"Attempting to extract text from DOCX: {file_path}")
    try:
        doc = Document(file_path)
        text = ""
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            text += para.text + "\n"
            
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        logger.debug(f"Successfully extracted {len(text)} characters from DOCX")
        return text
        
    except Exception as e:
        logger.error(f"Error in DOCX extraction: {str(e)}")
        raise

def extract_text_from_image(file_path):
    logger.debug(f"Attempting to extract text from image: {file_path}")
    try:
        image = Image.open(file_path)
        
        # Improve image quality for OCR
        image = image.convert('L')  # Convert to grayscale
        text = pytesseract.image_to_string(image)
        
        logger.debug(f"Successfully extracted {len(text)} characters from image")
        return text
        
    except Exception as e:
        logger.error(f"Error in image extraction: {str(e)}")
        raise

def clean_text_for_gpt(text: str) -> str:
    """Clean text before sending to GPT to avoid parsing issues."""
    # Remove special characters and normalize whitespace
    text = re.sub(r'[^\x00-\x7F]+', ' ', text)  # Remove non-ASCII characters
    text = re.sub(r'\s+', ' ', text)  # Normalize whitespace
    text = text.replace('`', "'")  # Replace backticks with single quotes
    text = text.replace('"', "'")  # Replace double quotes with single quotes
    return text.strip()

def validate_json(json_string: str) -> str:
    """Validate and clean JSON string."""
    try:
        # First try to parse it as is
        json.loads(json_string)
        return json_string
    except json.JSONDecodeError as e:
        logger.debug(f"Initial JSON parsing failed: {str(e)}")
        
        # Clean the string and try again
        cleaned = json_string
        # Remove any markdown code block markers
        cleaned = re.sub(r'```json\s*', '', cleaned)
        cleaned = re.sub(r'```\s*', '', cleaned)
        # Ensure proper quote usage
        cleaned = cleaned.replace("'", '"')
        
        try:
            # Validate the cleaned JSON
            json.loads(cleaned)
            return cleaned
        except json.JSONDecodeError:
            logger.error("Failed to clean JSON response")
            raise

def analyze_with_gpt(text: str):
    logger.debug("Analyzing text with GPT")
    try:
        # Clean the input text
        cleaned_text = clean_text_for_gpt(text)
        
        prompt = """
        Analyze this syllabus text and extract the following information in a structured format.
        Return ONLY valid JSON without any markdown formatting or explanation.
        The JSON should follow this structure:
        {
            "courseInfo": {
                "title": "",
                "description": "",
                "courseCode": ""
            },
            "gradeDistribution": {
                "weights": [{"category": "", "percentage": 0}],
                "scale": [{"grade": "", "minimum": 0}]
            },
            "policies": {
                "attendance": "",
                "lateWork": "",
                "examFormat": "",
                "homeworkFormat": "",
                "other": [{"title": "", "content": ""}]
            },
            "instructorInfo": {
                "instructors": [{"name": "", "email": "", "office": "", "officeHours": ""}],
                "tas": [{"name": "", "email": "", "officeHours": ""}]
            },
            "materials": [],
            "schedule": {
                "entries": [{"week": null, "date": "", "topic": "", "assignments": ""}]
            }
        }
        """

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": "You are a syllabus analyzer that extracts structured information from course syllabi. Return only valid JSON without any markdown formatting or additional text."
                },
                {
                    "role": "user",
                    "content": f"{prompt}\n\nSyllabus text:\n{cleaned_text}"
                }
            ],
            temperature=0.3  # Lower temperature for more consistent formatting
        )
        
        analyzed_content = response.choices[0].message.content
        logger.debug("GPT Analysis Result:")
        logger.debug("-" * 50)
        logger.debug(analyzed_content)
        logger.debug("-" * 50)
        
        # Validate and clean the JSON response
        validated_json = validate_json(analyzed_content)
        
        return validated_json
        
    except Exception as e:
        logger.error(f"Error in GPT analysis: {str(e)}")
        raise

@app.route('/extract-text', methods=['POST'])
@limiter.limit("12 per hour; 24 per day")
def extract_text():
    logger.debug("Received text extraction request")
    
    if 'file' not in request.files:
        logger.error("No file provided in request")
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    logger.debug(f"Processing file: {file.filename}")
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file.filename.split('.')[-1]}") as temp:
        try:
            file.save(temp.name)
            logger.debug(f"Saved uploaded file to temporary location: {temp.name}")
            
            file_extension = file.filename.split('.')[-1].lower()
            logger.debug(f"Detected file extension: {file_extension}")
            
            # Extract text based on file type
            if file_extension == 'pdf':
                text = extract_text_from_pdf(temp.name)
            elif file_extension == 'docx':
                text = extract_text_from_docx(temp.name)
            elif file_extension in ['png', 'jpg', 'jpeg']:
                text = extract_text_from_image(temp.name)
            else:
                logger.debug("Attempting to read as plain text file")
                with open(temp.name, 'r', encoding='utf-8') as f:
                    text = f.read()
            
            # Log the extracted text for debugging
            logger.debug(f"Extracted text preview: {text[:500]}...")
            
            # Analyze text with GPT
            analyzed_content = analyze_with_gpt(text)
            
            # Validate the analyzed content is proper JSON
            json.loads(analyzed_content)  # This will raise an error if invalid
            
            return jsonify({
                'text': text,
                'analyzed': analyzed_content,
                'filename': file.filename,
                'chars_extracted': len(text)
            })
            
        except json.JSONDecodeError as e:
            logger.error(f"JSON parsing error: {str(e)}")
            return jsonify({
                'error': f"Failed to parse GPT response: {str(e)}",
                'filename': file.filename
            }), 500
        except Exception as e:
            logger.error(f"Error processing file: {str(e)}")
            return jsonify({
                'error': str(e),
                'filename': file.filename
            }), 500
        finally:
            try:
                os.unlink(temp.name)
                logger.debug("Cleaned up temporary file")
            except Exception as e:
                logger.error(f"Error cleaning up temporary file: {str(e)}")

# Add error handler for rate limit exceeded
@app.errorhandler(429)
def ratelimit_handler(e):
    logger.warning(f"Rate limit exceeded: {e.description}")
    return jsonify({
        'error': 'Rate limit exceeded',
        'message': 'You have reached your usage limit. Please try again later.',
        'retry_after': e.retry_after if hasattr(e, 'retry_after') else None
    }), 429

if __name__ == '__main__':
    logger.info("Starting Flask server...")
    app.run(port=5000, debug=True) 